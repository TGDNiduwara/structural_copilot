"""
tools/word_tool.py
===================
Formal structural calculation report generator using python-docx.

Produces an engineering-styled Word document with a title block, assumptions,
design standards, tabulated support reactions and governing member forces,
and embedded SFD/BMD diagram images.

Requires: python-docx (`pip install python-docx`)
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from typing import List, Optional

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger("structural_copilot.word_tool")
logger.setLevel(logging.INFO)

NAVY = RGBColor(0x1F, 0x38, 0x64)
STEEL_BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x59, 0x59, 0x59)


class WordReporter:
    """Generates a formal structural design calculation report (.docx)."""

    def __init__(self):
        self.document: Optional[Document] = None

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #

    def generate_calculation_report(
        self,
        file_path: str,
        project_title: str,
        engineer_name: str,
        summary_text: str,
        member_df: pd.DataFrame,
        reactions_df: pd.DataFrame,
        diagram_paths: Optional[List[str]] = None,
        design_standards: Optional[List[str]] = None,
        assumptions: Optional[List[str]] = None,
    ) -> str:
        """Builds and saves the full calculation report; returns the file path."""
        self.document = Document()
        self._configure_base_styles()

        self._add_title_section(project_title, engineer_name)
        self._add_section_heading("1. Project Assumptions")
        self._add_bullet_list(
            assumptions or self._default_assumptions()
        )

        self._add_section_heading("2. Design Standards & Codes")
        self._add_bullet_list(
            design_standards or self._default_standards()
        )

        self._add_section_heading("3. Executive Summary")
        self._add_body_paragraph(summary_text)

        self._add_section_heading("4. Support Reactions")
        self._add_dataframe_table(reactions_df)

        self._add_section_heading("5. Governing Member Forces")
        self._add_dataframe_table(self._extract_governing_forces(member_df))

        if diagram_paths:
            self._add_section_heading("6. Shear Force & Bending Moment Diagrams")
            for path in diagram_paths:
                self._embed_diagram(path)

        self._add_footer_note()

        os.makedirs(os.path.dirname(os.path.abspath(file_path)) or ".", exist_ok=True)
        self.document.save(file_path)
        logger.info("Calculation report saved to %s", file_path)
        return file_path

    # ------------------------------------------------------------------ #
    # Style configuration
    # ------------------------------------------------------------------ #

    def _configure_base_styles(self) -> None:
        section = self.document.sections[0]
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

        normal = self.document.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)

        heading1 = self.document.styles["Heading 1"]
        heading1.font.name = "Calibri"
        heading1.font.size = Pt(14)
        heading1.font.bold = True
        heading1.font.color.rgb = NAVY

    # ------------------------------------------------------------------ #
    # Section builders
    # ------------------------------------------------------------------ #

    def _add_title_section(self, project_title: str, engineer_name: str) -> None:
        title_para = self.document.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("STRUCTURAL DESIGN CALCULATION REPORT")
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = NAVY

        subtitle = self.document.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle.add_run(project_title)
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = STEEL_BLUE
        sub_run.font.bold = True

        self._add_horizontal_rule()

        meta_table = self.document.add_table(rows=3, cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        meta_table.autofit = True

        meta_rows = [
            ("Project", project_title),
            ("Prepared By", engineer_name),
            ("Date Generated", datetime.now().strftime("%d %B %Y, %H:%M")),
        ]
        for row_idx, (label, value) in enumerate(meta_rows):
            label_cell = meta_table.rows[row_idx].cells[0]
            value_cell = meta_table.rows[row_idx].cells[1]
            # [FIX H6] Use helper to safely set text and access runs
            self._set_cell_text_safe(label_cell, str(label), bold=True, color=GREY)
            self._set_cell_text_simple(value_cell, str(value))
            for cell in (label_cell, value_cell):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        self.document.add_paragraph()  # spacer

    def _add_section_heading(self, text: str) -> None:
        heading = self.document.add_heading(text, level=1)
        heading.paragraph_format.space_before = Pt(18)
        heading.paragraph_format.space_after = Pt(6)

    def _add_body_paragraph(self, text: str) -> None:
        para = self.document.add_paragraph(text)
        para.paragraph_format.space_after = Pt(10)
        para.paragraph_format.line_spacing = 1.15

    def _add_bullet_list(self, items: List[str]) -> None:
        for item in items:
            para = self.document.add_paragraph(item, style="List Bullet")
            para.paragraph_format.space_after = Pt(3)

    def _add_dataframe_table(self, df: pd.DataFrame) -> None:
        if df is None or df.empty:
            empty_para = self.document.add_paragraph(
                "No data available for this section."
            )
            # [FIX H6] Safe run access
            if empty_para.runs:
                empty_para.runs[0].font.italic = True
                empty_para.runs[0].font.color.rgb = GREY
            else:
                run = empty_para.add_run("")
                run.font.italic = True
                run.font.color.rgb = GREY
            return

        table = self.document.add_table(rows=1, cols=len(df.columns))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name).replace("_", " ")
            for p in hdr_cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            self._shade_cell(hdr_cells[i], "2E75B6")

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, col_name in enumerate(df.columns):
                value = row[col_name]
                if isinstance(value, float):
                    text = f"{value:.2f}"
                else:
                    text = str(value)
                row_cells[i].text = text
                for p in row_cells[i].paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.document.add_paragraph()  # spacer after table

    def _embed_diagram(self, image_path: str) -> None:
        if not image_path or not os.path.isfile(image_path):
            logger.warning("Diagram image not found, skipping embed: %s", image_path)
            return

        caption = os.path.splitext(os.path.basename(image_path))[0].replace("_", " ").title()
        self.document.add_picture(image_path, width=Inches(6.3))
        last_paragraph = self.document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cap_para = self.document.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(f"Figure — {caption}")
        cap_run.font.size = Pt(9)
        cap_run.font.italic = True
        cap_run.font.color.rgb = GREY
        cap_para.paragraph_format.space_after = Pt(14)

    def _add_footer_note(self) -> None:
        self._add_horizontal_rule()
        note = self.document.add_paragraph()
        note.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = note.add_run(
            "This report was generated automatically by the Structural "
            "Multi-App Agent. Results should be independently verified by a "
            "licensed structural engineer prior to construction issuance."
        )
        run.font.size = Pt(8.5)
        run.font.italic = True
        run.font.color.rgb = GREY

    # ------------------------------------------------------------------ #
    # Low-level formatting helpers
    # ------------------------------------------------------------------ #

    def _add_horizontal_rule(self) -> None:
        para = self.document.add_paragraph()
        p_pr = para._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F3864")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)

    def _shade_cell(self, cell, hex_color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def _set_cell_text_safe(
        self, cell, text: str, bold: bool = False, color: RGBColor = None
    ) -> None:
        """[FIX H6] Safely sets cell text and applies formatting,
        ensuring at least one run exists before accessing runs[0]."""
        cell.text = text
        for paragraph in cell.paragraphs:
            if not paragraph.runs:
                # Ensure a run exists if text was set but produced no runs
                if paragraph.text:
                    # Re-setting text on the paragraph forces a run creation
                    continue
                else:
                    paragraph.add_run(text)
            for run in paragraph.runs:
                if bold:
                    run.font.bold = True
                if color:
                    run.font.color.rgb = color

    def _set_cell_text_simple(self, cell, text: str) -> None:
        """Sets cell text value without special formatting."""
        cell.text = text

    # ------------------------------------------------------------------ #
    # Default content
    # ------------------------------------------------------------------ #

    @staticmethod
    def _default_assumptions() -> List[str]:
        return [
            "All members are modeled as linear-elastic prismatic beam elements.",
            "Self-weight of structural steel is included via material density unless stated otherwise.",
            "Supports are idealized as indicated in the model (fixed / pinned / roller) with no soil-structure interaction considered.",
            "Loads are applied as static, non-dynamic actions; seismic and wind dynamic amplification are outside the scope of this report unless separately noted.",
            "Second-order (P-Delta) effects are neglected unless explicitly enabled in the analysis case.",
        ]

    @staticmethod
    def _default_standards() -> List[str]:
        return [
            "Eurocode 3 (EN 1993-1-1) — Design of steel structures, General rules.",
            "Eurocode 1 (EN 1991-1-1) — Actions on structures, Densities, self-weight, imposed loads.",
            "Eurocode 0 (EN 1990) — Basis of structural design, load combination factors.",
            "Autodesk Robot Structural Analysis Professional — Finite Element solver, linear static analysis.",
        ]

    @staticmethod
    def _extract_governing_forces(member_df: pd.DataFrame) -> pd.DataFrame:
        """Reduces a full member-forces export to one governing (max |MY|) row per bar."""
        if member_df is None or member_df.empty:
            return pd.DataFrame(columns=["Bar_ID", "Position_m", "FX_kN", "FZ_kN", "MY_kNm"])

        # [FIX M14] Check that required column exists
        if "MY_kNm" not in member_df.columns:
            logger.warning("member_df missing 'MY_kNm' column; returning empty governing forces.")
            return pd.DataFrame(columns=["Bar_ID", "Position_m", "FX_kN", "FZ_kN", "MY_kNm"])

        df = member_df.copy()
        df["_abs_my"] = df["MY_kNm"].abs()
        idx = df.groupby("Bar_ID")["_abs_my"].idxmax()
        governing = df.loc[idx].drop(columns="_abs_my").sort_values("Bar_ID").reset_index(drop=True)
        return governing
